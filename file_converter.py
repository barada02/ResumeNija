import io
import os
import re
from typing import Optional

# PDF handling
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from PyPDF2 import PdfReader

# DOCX handling
import docx

class FileConverter:
    """Converts different file formats to plain text for resume parsing"""
    
    @staticmethod
    def convert_to_text(file_content: bytes, file_type: str) -> str:
        """Convert file content to text based on file type
        
        Args:
            file_content: Binary content of the file
            file_type: Type of file ('pdf', 'docx', 'txt', 'md')
            
        Returns:
            Extracted text from the file
        """
        if file_type == 'pdf':
            return FileConverter._convert_pdf(file_content)
        elif file_type == 'docx':
            return FileConverter._convert_docx(file_content)
        elif file_type in ['txt', 'md']:
            # Text files can be decoded directly
            return file_content.decode('utf-8', errors='replace')
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    @staticmethod
    def _convert_pdf(file_content: bytes) -> str:
        """Extract text from PDF file
        
        Uses both pdfminer.six and PyPDF2 for better extraction results
        and falls back to the other if one fails.
        """
        text = ""
        
        # Try with pdfminer.six first (better for formatting)
        try:
            resource_manager = PDFResourceManager()
            fake_file_handle = io.StringIO()
            converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
            page_interpreter = PDFPageInterpreter(resource_manager, converter)
            
            for page in PDFPage.get_pages(io.BytesIO(file_content)):
                page_interpreter.process_page(page)
                
            text = fake_file_handle.getvalue()
            converter.close()
            fake_file_handle.close()
        except Exception as e:
            print(f"pdfminer.six extraction failed: {e}")
            
        # If pdfminer.six failed or returned empty text, try PyPDF2
        if not text.strip():
            try:
                pdf = PdfReader(io.BytesIO(file_content))
                text = ""
                for page in pdf.pages:
                    text += page.extract_text() + "\n"
            except Exception as e:
                print(f"PyPDF2 extraction failed: {e}")
        
        # Clean up the text
        text = FileConverter._clean_text(text)
        return text
    
    @staticmethod
    def _convert_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            full_text = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                full_text.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))
            
            text = '\n'.join(full_text)
            return FileConverter._clean_text(text)
        except Exception as e:
            print(f"DOCX extraction failed: {e}")
            return ""
    
    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean up extracted text"""
        # Replace multiple newlines with a single newline
        text = re.sub(r'\n\s*\n', '\n', text)
        
        # Replace multiple spaces with a single space
        text = re.sub(r'\s+', ' ', text)
        
        # Remove non-printable characters
        text = re.sub(r'[^\x20-\x7E\n]', '', text)
        
        return text.strip()
